"""Render Table 1 as a Word (.docx) document using python-docx."""
from __future__ import annotations
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_border(cell, **kwargs):
    """Set borders on a table cell. kwargs: top, bottom, left, right -> {sz, val, color}."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            border = tcBorders.find(qn(f"w:{edge}"))
            if border is None:
                border = OxmlElement(f"w:{edge}")
                tcBorders.append(border)
            for k, v in kwargs[edge].items():
                border.set(qn(f"w:{k}"), str(v))


def _design_note_text(design: str) -> str:
    if design == "RCT":
        return ("Per CONSORT 2010 (Moher BMJ 2010;340:c869) and Senn (Stat Med 1994;13:1715-26), "
                "baseline p-values are NOT reported for RCTs. Balance is assessed by SMD.")
    if design in ("Prospective cohort", "Retrospective cohort", "Cross-sectional", "Registry"):
        return ("Observational design. Unadjusted p-values reported as supportive only; "
                "interpret with multiple-testing caution. SMD is the primary balance metric "
                "(Austin PC, Stat Med 2009;28:3083): |SMD|>=0.1 small, >=0.2 meaningful.")
    if design == "Case-control":
        return ("Case-control design. Variables used for case selection may not be meaningfully "
                "compared. For matched designs, consider conditional logistic regression.")
    if design == "Single-arm prospective":
        return "Single-arm: descriptive only; between-group comparisons not applicable."
    return ("Design not specified. p-values are unadjusted and not corrected for multiple testing.")


def render_docx(rows: list[dict], meta: dict, out_path: str):
    doc = Document()

    # ---- Title ----
    title = doc.add_paragraph()
    run = title.add_run("Table 1. Baseline Characteristics")
    run.bold = True
    run.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub_run = sub.add_run(
        f"{meta['input_name']} - grouped by {meta['group_var']} - "
        f"Study design: {meta['study_design']} - {meta['generated_at']}"
    )
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # ---- Design note ----
    note = doc.add_paragraph()
    nr = note.add_run(_design_note_text(meta["study_design"]))
    nr.italic = True
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x30, 0x30, 0x30)

    # ---- Table ----
    group_levels = meta["group_levels"]
    group_n = meta["group_n"]
    n_grp = len(group_levels)

    # Columns: Variable | Group1 | Group2 | ... | Missing | Test | p | SMD
    n_cols = 1 + n_grp + 4
    table = doc.add_table(rows=1, cols=n_cols)
    table.autofit = True
    table.style = "Light Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Variable"
    for i, g in enumerate(group_levels):
        hdr_cells[1 + i].text = f"{g}\n(n={group_n.get(g, 0)})"
    hdr_cells[1 + n_grp].text = "Missing"
    hdr_cells[2 + n_grp].text = "Test"
    hdr_cells[3 + n_grp].text = "p"
    hdr_cells[4 + n_grp].text = "SMD"

    for c in hdr_cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        _set_cell_border(c, bottom={"sz": 12, "val": "single", "color": "000000"})

    # Body rows
    for row in rows:
        rcells = table.add_row().cells
        # Variable name
        name = row["name"]
        if row["kind"] == "level":
            name = "    " + name.lstrip()
        annot = row.get("annotation", "")
        if annot and row["kind"] != "level":
            name_text = f"{name}  ({annot})"
        else:
            name_text = name
        rcells[0].text = name_text
        for i, g in enumerate(group_levels):
            rcells[1 + i].text = row["per_group"].get(g, "")
        rcells[1 + n_grp].text = row["missing"]
        rcells[2 + n_grp].text = row["test"]
        rcells[3 + n_grp].text = row["p_str"]
        smd_text = row["smd_str"]
        if row.get("smd_lbl"):
            smd_text += f" ({row['smd_lbl']})"
        rcells[4 + n_grp].text = smd_text

        # Style
        for ci, c in enumerate(rcells):
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                if ci > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Bold header for categorical_header
        if row["kind"] == "categorical_header":
            for p in rcells[0].paragraphs:
                for r in p.runs:
                    r.bold = True

        # SMD color
        band = row.get("smd_band", "")
        if band:
            color_map = {
                "ok": RGBColor(0x15, 0x80, 0x3d),
                "small": RGBColor(0xca, 0x8a, 0x04),
                "meaningful": RGBColor(0xb9, 0x1c, 0x1c),
            }
            color = color_map.get(band)
            if color:
                for p in rcells[4 + n_grp].paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = color
                        r.bold = True

    # ---- Footer notes ----
    doc.add_paragraph()
    f1 = doc.add_paragraph()
    fr = f1.add_run(
        "Methods: Continuous variables summarized as mean +/- SD or median [IQR] based on "
        "skewness (|skew|>1) and known-skewed lab variables. Tests: Welch t / Mann-Whitney U "
        "(2-group); one-way ANOVA / Kruskal-Wallis (>=3-group). Categorical: chi-square (expected >=5), "
        "Fisher's exact (2x2 with small cells), Monte Carlo chi-square (rxc with small cells). "
        "SMD: standardized differences; for >=3 groups, the maximum pairwise SMD is shown."
    )
    fr.font.size = Pt(8)
    fr.italic = True

    f2 = doc.add_paragraph()
    fr2 = f2.add_run(
        f"Source: {meta['input_name']} (SHA-256 {meta['input_sha256']}). "
        f"Generated by clinical-table1 on {meta['generated_at']}."
    )
    fr2.font.size = Pt(8)
    fr2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.save(out_path)
